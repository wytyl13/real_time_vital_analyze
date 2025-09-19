#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
@Time    : 2025/09/02 16:01
@Author  : weiyutao
@File    : structured_docx_extractor.py

开发到一半，转而使用大模型实现这个功能！
"""
from docx import Document
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from typing import Dict, Optional, Any, List, Union
import json
import re


class StructuredDocxExtractor:
    """基于文档结构的DocX字段提取器"""
    
    def __init__(self):
        # 字段映射 - 用于匹配不同的字段名称变体
        self.extract_field = [
            {"from": "姓  名", "to": "name", "type": "paragraph"},
            {"from": "性  别", "to": "gender", "type": "paragraph"},
            {"from": "年  龄", "to": "age", "type": "paragraph"},
            {"from": "档案号", "to": "file_number", "type": "paragraph"},
            {"from": "评估人", "to": "evaluator", "type": "paragraph"},
            {"from": "评估医生", "to": "evaluation_doctor", "type": "paragraph"},
            {"from": "审核人", "to": "reviewer", "type": "paragraph"},
            {"from": "服务对象", "to": "service_object", "type": "paragraph"},
            {"from": "服务对象家属", "to": "service_object_family", "type": "paragraph"},
            {"from": "评估日期", "to": "evalution_date", "type": "paragraph"},
            {"from": "走失风险评估总分", "to": "evalution_date", "type": "table"},
            {"from": "走失风险评估分级", "to": "evalution_date", "type": "table"},
        ]
        # 构建字段映射字典
        self.field_mappings = {}
        for item in self.extract_field:
            original_field = item["from"]
            self.field_mappings[original_field] = [original_field]
        print(self.field_mappings)


    def extract_from_document(self, file_path: str) -> List[Dict[str, str]]:
        """从docx文档中提取字段"""
        try:
            doc = Document(file_path)
            results = {}
            
            # 1. 从段落中提取（处理表单式布局）
            paragraph_data = self._extract_from_paragraphs(doc.paragraphs)
            results.update(paragraph_data)
            
            # 2. 从表格中提取
            for table in doc.tables:
                table_data = self._extract_from_table(table)
                results.update(table_data)
            
            # 3. 后处理和验证
            results = self._post_process_results(results)
            
            # 转换为新的返回格式
            final_results = []
            for item in self.extract_field:
                from_field = item["from"]
                to_field = item["to"]
                value = results.get(from_field, "")
                final_results.append({
                    "from": to_field,
                    "to": value
                })
            
            return final_results
            
        except Exception as e:
            return [{"error": str(e)}]


    def _extract_from_paragraphs(self, paragraphs: List[Paragraph]) -> Dict[str, str]:
        """从段落中提取字段 - 处理表单式布局"""
        results = {}
        
        for paragraph in paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            print(text)
                
            # 方法1: 处理单行多字段 (如: 姓名：张三  性别：男  年龄：25)
            results.update(self._extract_inline_fields(text))
            
            # 方法2: 处理单独行的字段 (如: 姓名：张三)
            field_result = self._extract_single_field(text)
            if field_result:
                results.update(field_result)
        print(results)
        return results


    def _extract_from_table(self, table: Table) -> Dict[str, str]:
        """从表格中提取字段"""
        results = {}
        
        # 遍历表格的每一行
        for row_idx, row in enumerate(table.rows):
            row_cells = [cell.text.strip() for cell in row.cells]
            
            # 方法1: 两列表格 (字段名 | 字段值)
            if len(row_cells) >= 2:
                field_name = row_cells[0]
                field_value = row_cells[1]
                
                standardized_field = self._standardize_field_name(field_name)
                if standardized_field and field_value:
                    results[standardized_field] = self._clean_field_value(field_value)
            
            # 方法2: 单行多列表格处理
            row_text = ' '.join(row_cells)
            results.update(self._extract_inline_fields(row_text))
        
        return results


    def _extract_inline_fields(self, text: str) -> Dict[str, str]:
        """提取一行中的多个字段"""
        results = {}
        
        # 修改正则模式，遇到空格后停止提取
        patterns = [
            # 支持字段名中包含空格的模式，如 "姓  名：值"
            r'([^：:]{1,10}?)\s*[：:]\s*([^\s：:]+?)(?=\s+[^：:]{1,10}?\s*[：:]|\s*$)',
            # 支持下划线填充的字段
            r'([^：:_]{1,10}?)\s*[：:]\s*_*([^_\s]+?)_*(?=\s+[^：:_]{1,10}?\s*[：:]|\s*$)'
        ]
        
        for pattern in patterns:
            matches = re.findall(pattern, text)
            for field_name, field_value in matches:
                field_name = field_name.strip()
                field_value = field_value.strip()
                
                standardized_field = self._standardize_field_name(field_name)
                if standardized_field and field_value and field_value != '_':
                    results[standardized_field] = self._clean_field_value(field_value)
        
        return results


    def _extract_single_field(self, text: str) -> Optional[Dict[str, str]]:
        """提取单个字段"""
        patterns = [
            r'^([^：:\s]{1,10})\s*[：:]\s*(.+)$',
            r'^([^：:\s]{1,10})\s*[：:]\s*_*(.+?)_*$'
        ]
        
        for pattern in patterns:
            match = re.match(pattern, text.strip())
            if match:
                field_name = match.group(1).strip()
                field_value = match.group(2).strip()
                
                standardized_field = self._standardize_field_name(field_name)
                if standardized_field and field_value and field_value != '_':
                    return {standardized_field: self._clean_field_value(field_value)}
        
        return None


    def _standardize_field_name(self, field_name: str) -> Optional[str]:
        field_name = field_name.strip()
        # print(f"尝试匹配字段: '{field_name}'")
        
        for target_field in self.field_mappings.keys():
            # print(f"  与目标字段比较: '{target_field}'")
            
            # 策略1: 精确匹配
            if field_name == target_field:
                # print(f"    精确匹配成功!")
                return target_field
            
            # 策略2: 去除所有空格后匹配
            target_clean = re.sub(r'\s+', '', target_field)
            input_clean = re.sub(r'\s+', '', field_name)
            # print(f"    去空格比较: '{input_clean}' vs '{target_clean}'")
            
            if input_clean == target_clean:
                # print(f"    去空格匹配成功!")
                return target_field
            
            # 策略3: 包含匹配
            if target_clean in input_clean:
                # print(f"    包含匹配成功!")
                return target_field
        
        # print(f"字段匹配失败...")
        return None


    def _clean_field_value(self, value: str) -> str:
        """清理字段值"""
        value = value.strip()
        # 如果包含空格，只取第一个单词
        if ' ' in value:
            value = value.split(' ')[0]
        value = re.sub(r'^_+|_+$', '', value)  # 移除首尾下划线
        value = re.sub(r'\s+', ' ', value)     # 规范化空格
        return value


    def _post_process_results(self, results: Dict[str, str]) -> Dict[str, str]:
        """后处理结果"""
        # 日期格式标准化
        if '评估日期' in results:
            date_value = results['评估日期']
            # 标准化日期格式
            date_patterns = [
                (r'(\d{4})年(\d{1,2})月(\d{1,2})日', r'\1-\2-\3'),
                (r'(\d{4})/(\d{1,2})/(\d{1,2})', r'\1-\2-\3'),
                (r'(\d{4})\.(\d{1,2})\.(\d{1,2})', r'\1-\2-\3')
            ]
            for pattern, replacement in date_patterns:
                date_value = re.sub(pattern, replacement, date_value)
            results['评估日期'] = date_value
        
        # 年龄数据验证
        if '年龄' in results:
            age_match = re.search(r'\d+', results['年龄'])
            if age_match:
                results['年龄'] = age_match.group()
        
        return results


    def extract_with_position_info(self, file_path: str) -> Dict[str, Any]:
        """提取字段并包含位置信息（用于调试）"""
        doc = Document(file_path)
        results = {'fields': {}, 'debug_info': []}
        
        # 记录段落信息
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():
                results['debug_info'].append({
                    'type': 'paragraph',
                    'index': i,
                    'text': paragraph.text.strip(),
                    'extracted_fields': self._extract_inline_fields(paragraph.text)
                })
        
        # 记录表格信息
        for table_idx, table in enumerate(doc.tables):
            table_info = {
                'type': 'table',
                'index': table_idx,
                'rows': len(table.rows),
                'columns': len(table.columns) if table.rows else 0,
                'extracted_fields': {}
            }
            
            for row_idx, row in enumerate(table.rows):
                row_cells = [cell.text.strip() for cell in row.cells]
                table_info['extracted_fields'][f'row_{row_idx}'] = {
                    'cells': row_cells,
                    'extracted': self._extract_inline_fields(' '.join(row_cells))
                }
            
            results['debug_info'].append(table_info)
        
        # 汇总所有提取的字段
        all_extracted = self.extract_from_document(file_path)
        results['fields'] = all_extracted
        
        return results


    def print_extraction_results(self, results: Dict[str, Any], show_debug: bool = False):
        """打印提取结果"""
        print("=" * 60)
        print("📋 结构化字段提取结果")
        print("=" * 60)
        
        if 'error' in results:
            print(f"❌ 提取失败: {results['error']}")
            return
        
        # 打印主要字段结果
        field_order = ['姓名', '性别', '年龄', '档案号', '评估人', '评估医生', '审核人', '服务对象', '服务对象家属', '评估日期']
        
        print("🎯 提取的字段:")
        for field_name in field_order:
            value = results.get(field_name)
            status = "✅" if value else "❌"
            print(f"  {status} {field_name:12}: {value or '未找到'}")
        
        # 显示其他提取到的字段
        other_fields = {k: v for k, v in results.items() 
                       if not k.startswith('_') and k not in field_order and v}
        if other_fields:
            print("\n🔍 其他提取的字段:")
            for field_name, value in other_fields.items():
                print(f"  ✅ {field_name:12}: {value}")
        
        # 打印统计信息
        if '_extraction_info' in results:
            info = results['_extraction_info']
            print(f"\n📊 提取统计:")
            print(f"   成功字段: {info.get('extracted_fields', 0)} / {info.get('total_target_fields', 0)}")
            print(f"   文档结构: {info.get('document_paragraphs', 0)} 个段落, {info.get('document_tables', 0)} 个表格")
            print(f"   提取方法: {', '.join(info.get('extraction_methods', []))}")
        
        # 调试信息
        if show_debug and 'debug_info' in results:
            print(f"\n🔧 调试信息:")
            for item in results['debug_info']:
                if item['type'] == 'paragraph':
                    print(f"   段落 {item['index']}: {item['text'][:50]}...")
                    if item['extracted_fields']:
                        print(f"      提取: {item['extracted_fields']}")
                elif item['type'] == 'table':
                    print(f"   表格 {item['index']}: {item['rows']}行 x {item['columns']}列")


if __name__ == '__main__':
    extractor = StructuredDocxExtractor()
    
    result = extractor.extract_from_document(file_path='/work/ai/real_time_vital_analyze/api/source/100.docx')
    print(result)