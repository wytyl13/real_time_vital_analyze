#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
@Time    : 2025/09/02 15:10
@Author  : weiyutao
@File    : test_server.py
"""


#!/usr/bin/env python3
"""
OnlyOffice在线文档编辑器 - FastAPI版本，支持URL传参和回调保存
"""

from fastapi import FastAPI, Request, HTTPException, Query
from fastapi.responses import HTMLResponse, JSONResponse, FileResponse
from pathlib import Path
import hashlib
import json
import time
import os
import requests
import tempfile
import urllib.parse
from urllib.parse import urlparse
import mimetypes
import logging
from typing import Optional, List, Dict, Any
from datetime import datetime




#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
DocX文档字段提取工具 - 使用python-docx结构化提取
"""

from docx import Document
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from typing import Dict, Optional, Any, List, Union
import json
import re




from pydantic import BaseModel

class EditDocumentRequest(BaseModel):
    url: str
    filename: Optional[str] = None
    replace_information: Optional[str] = None # JSON字符串格式
    extract_information: Optional[str] = None # JSON字符串格式




class OnlyOfficeEditor:
    """OnlyOffice文档编辑器类"""
    
    def __init__(self, onlyoffice_server: str, save_api_url: str, jwt_secret: str = ""):
        self.onlyoffice_server = onlyoffice_server
        self.save_api_url = save_api_url
        self.jwt_secret = jwt_secret
        self.temp_dir = tempfile.gettempdir()
        self.logger = logging.getLogger(self.__class__.__name__)
        
        
        # 文档类型映射
        self.type_mapping = {
            'doc': 'word', 'docx': 'word', 
            'xls': 'cell', 'xlsx': 'cell',
            'ppt': 'slide', 'pptx': 'slide'
        }


    def replace_document_content(self, file_path: str, replace_rules: List[Dict[str, str]]) -> str:
        """
        替换docx文档中的内容
        
        Args:
            file_path: 文档文件路径
            replace_rules: 替换规则列表，格式为 [{"from": "原文本", "to": "新文本"}]
        
        Returns:
            str: 替换后的文件路径
        """
        try:
            if not replace_rules:
                return file_path
            
            # 检查文件是否为docx格式
            if not file_path.lower().endswith(('.docx', '.doc')):
                self.logger.info(f"文件 {file_path} 不是Word文档，跳过内容替换")
                return file_path
            
            self.logger.info(f"开始替换文档内容，规则数量: {len(replace_rules)}")
            
            # 加载文档
            document = Document(file_path)
            
            # 替换段落中的文本
            self._replace_in_paragraphs(document.paragraphs, replace_rules)
            
            # 替换表格中的文本
            for table in document.tables:
                self._replace_in_table(table, replace_rules)
            
            # 替换页眉页脚中的文本
            for section in document.sections:
                # 页眉
                if section.header:
                    self._replace_in_paragraphs(section.header.paragraphs, replace_rules)
                    for table in section.header.tables:
                        self._replace_in_table(table, replace_rules)
                
                # 页脚
                if section.footer:
                    self._replace_in_paragraphs(section.footer.paragraphs, replace_rules)
                    for table in section.footer.tables:
                        self._replace_in_table(table, replace_rules)
            
            # 生成新的文件路径
            base_name = os.path.splitext(os.path.basename(file_path))[0]
            new_file_path = os.path.join(
                self.temp_dir, 
                f"replaced_{int(time.time())}_{base_name}.docx"
            )
            
            # 保存修改后的文档
            document.save(new_file_path)
            
            self.logger.info(f"文档内容替换完成，保存到: {new_file_path}")
            
            # 删除原临时文件
            try:
                os.remove(file_path)
            except Exception as e:
                self.logger.warning(f"删除原临时文件失败: {e}")
            
            return new_file_path
            
        except Exception as e:
            self.logger.error(f"替换文档内容失败: {e}")
            # 如果替换失败，返回原文件路径
            return file_path
    
    
    def _replace_in_paragraphs(self, paragraphs: List[Paragraph], replace_rules: List[Dict[str, str]]):
        """在段落中进行文本替换，保持原有格式"""
        for paragraph in paragraphs:
            if not paragraph.text.strip():  # 跳过空段落
                continue
            for rule in replace_rules:
                from_text = rule.get('from', '')
                to_text = rule.get('to', '')
                
                if not from_text:
                    continue
                
                # 检查段落中是否包含要替换的文本
                import re
                pattern = r'\b' + re.escape(from_text) + r'\b'
                if re.search(pattern, paragraph.text):
                    self.logger.info(f"段落替换: '{from_text}' -> '{to_text}'")
                    
                    # 方法1: 先尝试简单的run级别替换
                    replaced = False
                    for run in paragraph.runs:
                        if re.search(pattern, run.text):
                            run.text = re.sub(pattern, to_text, run.text)
                            replaced = True
                    
                    # 方法2: 如果简单替换没有成功（可能是跨run的文本），使用复杂替换
                    if not replaced and re.search(pattern, paragraph.text):
                        self._replace_across_runs(paragraph, from_text, to_text)
    
    
    def _replace_across_runs_debug(self, paragraph: Paragraph, from_text: str, to_text: str):
        """处理跨run的文本替换，尽可能保持格式 - 调试版本"""
        import re
        full_text = paragraph.text
        pattern = r'\b' + re.escape(from_text) + r'\b'
        match = re.search(pattern, full_text)
        if not match:
            self.logger.warning(f"未找到匹配的文本: '{from_text}' 在 '{full_text}' 中")
            return
        
        # 找到替换位置
        start_pos = match.start()
        end_pos = start_pos + len(from_text)  # 先保持你原来的逻辑
        
        # 详细调试信息
        print(f"=== 跨run替换调试 ===")
        print(f"from_text: '{from_text}'")
        print(f"to_text: '{to_text}'")
        print(f"full_text: '{full_text}'")
        print(f"匹配位置: {start_pos}-{end_pos}")
        print(f"段落runs数量: {len(paragraph.runs)}")
        
        # 构建新的runs列表
        new_runs = []
        current_pos = 0
        
        # 分析每个run的位置关系
        for i, run in enumerate(paragraph.runs):
            run_start = current_pos
            run_end = current_pos + len(run.text)
            print(f"Run {i}: '{run.text}' 位置:{run_start}-{run_end}")
            current_pos = run_end
        
        # 重新开始处理
        current_pos = 0
        replacement_added = False
        
        for run in paragraph.runs:
            run_start = current_pos
            run_end = current_pos + len(run.text)
            
            print(f"处理run: '{run.text}' (位置:{run_start}-{run_end})")
            
            if run_end <= start_pos:
                # run完全在替换文本之前
                new_runs.append((run.text, run))
                print(f"  -> 添加前置run: '{run.text}'")
                
            elif run_start >= end_pos:
                # run完全在替换文本之后
                if not replacement_added:
                    new_runs.append((to_text, run))
                    replacement_added = True
                    print(f"  -> 添加替换文本: '{to_text}'")
                new_runs.append((run.text, run))
                print(f"  -> 添加后置run: '{run.text}'")
                
            elif run_start <= start_pos and run_end >= end_pos:
                # 替换文本完全在这个run内 - 这是最常见的情况
                print(f"  -> 替换文本完全在run内")
                
                before = run.text[:start_pos - run_start]
                after = run.text[end_pos - run_start:]
                
                print(f"     before: '{before}'")
                print(f"     after: '{after}'")
                
                if before:
                    new_runs.append((before, run))
                    print(f"  -> 添加before: '{before}'")
                
                new_runs.append((to_text, run))  # 关键：无条件添加替换文本
                replacement_added = True
                print(f"  -> 添加替换文本: '{to_text}'")
                
                if after:
                    new_runs.append((after, run))
                    print(f"  -> 添加after: '{after}'")
                    
            elif start_pos <= run_start and end_pos >= run_end:
                # 这个run完全被替换文本覆盖，跳过
                print(f"  -> run被完全覆盖，跳过")
                if not replacement_added:
                    new_runs.append((to_text, run))
                    replacement_added = True
                    print(f"  -> 添加替换文本: '{to_text}'")
                    
            elif run_start < start_pos < run_end:
                # run的后半部分需要被替换
                print(f"  -> run后半部分被替换")
                before = run.text[:start_pos - run_start]
                if before:
                    new_runs.append((before, run))
                    print(f"  -> 添加before: '{before}'")
                # 替换文本将在后续处理
                if not replacement_added:
                    new_runs.append((to_text, run))
                    replacement_added = True
                    print(f"  -> 添加替换文本: '{to_text}'")
                    
            elif run_start < end_pos < run_end:
                # run的前半部分需要被替换
                print(f"  -> run前半部分被替换")
                if not replacement_added:
                    new_runs.append((to_text, run))
                    replacement_added = True
                    print(f"  -> 添加替换文本: '{to_text}'")
                after = run.text[end_pos - run_start:]
                if after:
                    new_runs.append((after, run))
                    print(f"  -> 添加after: '{after}'")
            
            current_pos = run_end
        
        # 检查是否添加了替换文本
        if not replacement_added:
            self.logger.warning("替换文本未被添加！作为兜底添加")
            template_run = paragraph.runs[0] if paragraph.runs else None
            new_runs.append((to_text, template_run))
        
        # 打印最终的new_runs
        print("=== 最终new_runs ===")
        for i, (text, run) in enumerate(new_runs):
            print(f"  {i}: '{text}'")
        
        # 清除原有runs并添加新的
        paragraph.clear()
        for text, template_run in new_runs:
            if text or text == "":  # 允许空字符串，但不允许None
                new_run = paragraph.add_run(text)
                # 复制格式
                if template_run:
                    self._copy_run_format(template_run, new_run)
                print(f"  创建新run: '{text}'")
        
        # 最终验证
        final_text = paragraph.text
        print(f"=== 替换结果 ===")
        print(f"最终段落文本: '{final_text}'")
        print(f"期望包含: '{to_text}'")
        print(f"替换成功: {to_text in final_text}")
        print("=== 调试结束 ===")
    
    
    def _replace_across_runs(self, paragraph: Paragraph, from_text: str, to_text: str):
        """处理跨run的文本替换，尽可能保持格式 - 改进格式版本"""
        import re
        full_text = paragraph.text
        pattern = r'\b' + re.escape(from_text) + r'\b'
        match = re.search(pattern, full_text)
        if not match:
            return
        
        # 保存原段落格式
        original_paragraph = paragraph
        
        # 找到替换位置
        start_pos = match.start()
        end_pos = start_pos + len(from_text)
        
        # 构建新的runs列表
        new_runs = []
        current_pos = 0
        replacement_added = False
        
        for run in paragraph.runs:
            run_start = current_pos
            run_end = current_pos + len(run.text)
            
            if run_end <= start_pos:
                # run完全在替换文本之前
                if run.text:
                    new_runs.append((run.text, run))
                    
            elif run_start >= end_pos:
                # run完全在替换文本之后
                if not replacement_added:
                    # 使用最近的有效run作为格式模板
                    template_run = new_runs[-1][1] if new_runs else run
                    new_runs.append((to_text, template_run))
                    replacement_added = True
                if run.text:
                    new_runs.append((run.text, run))
                    
            elif run_start <= start_pos and run_end >= end_pos:
                # 替换文本完全在这个run内 - 最常见情况
                before = run.text[:start_pos - run_start]
                after = run.text[end_pos - run_start:]
                
                if before:
                    new_runs.append((before, run))
                new_runs.append((to_text, run))  # 使用原run的格式
                replacement_added = True
                if after:
                    new_runs.append((after, run))
                    
            elif start_pos <= run_start and end_pos >= run_end:
                # 这个run完全被替换文本覆盖
                if not replacement_added:
                    new_runs.append((to_text, run))
                    replacement_added = True
                    
            elif run_start < start_pos < run_end:
                # run的后半部分需要被替换
                before = run.text[:start_pos - run_start]
                if before:
                    new_runs.append((before, run))
                if not replacement_added:
                    new_runs.append((to_text, run))
                    replacement_added = True
                    
            elif run_start < end_pos < run_end:
                # run的前半部分需要被替换
                if not replacement_added:
                    new_runs.append((to_text, run))
                    replacement_added = True
                after = run.text[end_pos - run_start:]
                if after:
                    new_runs.append((after, run))
            
            current_pos = run_end
        
        # 如果还没有添加替换文本
        if not replacement_added:
            template_run = paragraph.runs[0] if paragraph.runs else None
            new_runs.append((to_text, template_run))
        
        # 清除原有runs并添加新的
        paragraph.clear()
        
        # 重建runs并复制格式
        for text, template_run in new_runs:
            if text or text == "":
                new_run = paragraph.add_run(text)
                # 使用改进的格式复制
                if template_run:
                    self._copy_run_format(template_run, new_run)
        
        # 复制段落级别的格式
        # 注意：这里我们无法直接复制，因为段落已经被修改
        # 但可以确保段落的基本属性保持一致
        self.logger.debug(f"替换完成: '{from_text}' -> '{to_text}'")
    
    
    def _copy_run_format(self, source_run, target_run):
        """复制run的格式 - 完整版本"""
        try:
            if not source_run or not target_run:
                return
                
            source_font = source_run.font
            target_font = target_run.font
            
            # 复制字体基本属性
            if source_font.bold is not None:
                target_font.bold = source_font.bold
            if source_font.italic is not None:
                target_font.italic = source_font.italic
            if source_font.underline is not None:
                target_font.underline = source_font.underline
            if source_font.name is not None:
                target_font.name = source_font.name
            if source_font.size is not None:
                target_font.size = source_font.size
                
            # 复制字体颜色
            try:
                if source_font.color.rgb is not None:
                    target_font.color.rgb = source_font.color.rgb
            except:
                pass
                
            try:
                if source_font.color.theme_color is not None:
                    target_font.color.theme_color = source_font.color.theme_color
            except:
                pass
                
            # 复制高亮颜色
            try:
                if source_font.highlight_color is not None:
                    target_font.highlight_color = source_font.highlight_color
            except:
                pass
                
            # 复制上标下标
            try:
                if source_font.subscript is not None:
                    target_font.subscript = source_font.subscript
                if source_font.superscript is not None:
                    target_font.superscript = source_font.superscript
            except:
                pass
                
            # 复制删除线和双删除线
            try:
                if hasattr(source_font, 'strike') and source_font.strike is not None:
                    target_font.strike = source_font.strike
                if hasattr(source_font, 'double_strike') and source_font.double_strike is not None:
                    target_font.double_strike = source_font.double_strike
            except:
                pass
                
            # 复制字符间距
            try:
                if hasattr(source_font, 'character_spacing') and source_font.character_spacing is not None:
                    target_font.character_spacing = source_font.character_spacing
            except:
                pass
            
            self.logger.debug(f"格式复制完成: {source_font.name}, {source_font.size}, bold:{source_font.bold}")
            
        except Exception as e:
            self.logger.warning(f"复制格式时出现警告: {e}")
    
    
    def _copy_run_format_simple(self, source_run, target_run):
        """复制run的格式"""
        try:
            # 复制字体设置
            if source_run.font.bold is not None:
                target_run.font.bold = source_run.font.bold
            if source_run.font.italic is not None:
                target_run.font.italic = source_run.font.italic
            if source_run.font.underline is not None:
                target_run.font.underline = source_run.font.underline
            if source_run.font.name is not None:
                target_run.font.name = source_run.font.name
            if source_run.font.size is not None:
                target_run.font.size = source_run.font.size
            if source_run.font.color.rgb is not None:
                target_run.font.color.rgb = source_run.font.color.rgb
            if source_run.font.highlight_color is not None:
                target_run.font.highlight_color = source_run.font.highlight_color
        except Exception as e:
            self.logger.warning(f"复制格式时出现警告: {e}")
    
    
    def _replace_in_table(self, table: Table, replace_rules: List[Dict[str, str]]):
        """在表格中进行文本替换"""
        for row in table.rows:
            for cell in row.cells:
                # 替换单元格段落中的文本
                self._replace_in_paragraphs(cell.paragraphs, replace_rules)
                
                # 处理嵌套表格
                for nested_table in cell.tables:
                    self._replace_in_table(nested_table, replace_rules)


    def download_file_from_url(self, url: str, filename: Optional[str] = None) -> tuple[str, str]:
        """从URL下载文件到临时目录"""
        try:
            response = requests.get(url, stream=True, timeout=30)
            response.raise_for_status()
            
            # 如果没有提供filename，尝试从响应中获取
            if not filename:
                # 首先尝试从Content-Disposition头获取
                content_disposition = response.headers.get('content-disposition')
                if content_disposition:
                    import re
                    match = re.search(r'filename[*]?=([^;]+)', content_disposition)
                    if match:
                        filename = match.group(1).strip('"\'')
                
                # 如果还没有filename，从URL路径获取
                if not filename:
                    parsed_url = urlparse(url)
                    filename = os.path.basename(parsed_url.path)
                
                # 如果仍然没有filename，使用默认名称
                if not filename or '.' not in filename:
                    content_type = response.headers.get('content-type', '')
                    if 'wordprocessing' in content_type:
                        filename = 'document.docx'
                    elif 'spreadsheet' in content_type:
                        filename = 'document.xlsx'
                    elif 'presentation' in content_type:
                        filename = 'document.pptx'
                    else:
                        filename = 'document.docx'
            
            # 确保filename有正确的扩展名
            if '.' not in filename:
                filename += '.docx'
            
            # 生成临时文件路径
            temp_path = os.path.join(self.temp_dir, f"onlyoffice_{int(time.time())}_{filename}")
            
            # 保存文件
            with open(temp_path, 'wb') as f:
                for chunk in response.iter_content(chunk_size=8192):
                    if chunk:
                        f.write(chunk)
            
            return temp_path, filename
            
        except Exception as e:
            raise Exception(f"下载文件失败: {str(e)}")


    def upload_file_to_api(self, file_content: bytes, filename: str) -> dict:
            """上传文件到保存API"""
            try:
                files = {
                    'file': (filename, file_content, 'application/octet-stream')
                }
                data = {
                    'filename': filename
                }
                
                response = requests.post(self.save_api_url, files=files, data=data, timeout=30)
                response.raise_for_status()
                
                return response.json()
            except Exception as e:
                raise Exception(f"上传文件到API失败: {str(e)}")


    def cleanup_temp_files(self):
        """清理临时文件"""
        try:
            for filename in os.listdir(self.temp_dir):
                if filename.startswith(('onlyoffice_', 'replaced_')) and os.path.isfile(os.path.join(self.temp_dir, filename)):
                    file_path = os.path.join(self.temp_dir, filename)
                    # 删除超过1小时的临时文件
                    if time.time() - os.path.getctime(file_path) > 3600:
                        os.remove(file_path)
                        self.logger.info(f"清理临时文件: {file_path}")
        except Exception as e:
            self.logger.error(f"清理临时文件出错: {e}")


    def generate_onlyoffice_config(self, file_path: str, filename: str, document_url: str) -> dict:
        """生成OnlyOffice配置"""
        file_ext = filename.lower().split('.')[-1]
        document_type = self.type_mapping.get(file_ext, 'word')
        
        # 生成文档密钥
        file_stat = os.stat(file_path)
        doc_key = hashlib.md5(f"{document_url}-{file_stat.st_mtime}".encode()).hexdigest()
        
        # 文件服务URL
        local_file_url = f"http://ai.shunxikj.com:8002/serve_temp_file?path={urllib.parse.quote(file_path)}"
        
        config = {
            "documentType": document_type,
            "document": {
                "fileType": file_ext,
                "key": doc_key,
                "title": filename,
                "url": local_file_url,
                "permissions": {
                    "comment": True,
                    "copy": True,
                    "download": True,
                    "edit": True,
                    "fillForms": True,              # 表单填写权限
                    "modifyFilter": True,           # 修改筛选器权限
                    "modifyContentControl": True,   # 修改内容控件权限
                    "review": True,
                    "print": True,
                    "changeHistory": True,          # 修改历史权限
                    "rename": True                  # 重命名权限
                }
            },
            "editorConfig": {
                "lang": "zh-CN",
                "mode": "edit",
                "callbackUrl": f"http://ai.shunxikj.com:8002/callback?key={doc_key}&filename={urllib.parse.quote(filename)}&original_url={urllib.parse.quote(document_url)}",
                "user": {
                    "id": "user-1",
                    "name": "编辑用户"
                },
                "customization": {
                    "leftMenu": False,
                    "toolbar": True,
                    "autosave": True,
                    "forcesave": False,
                    "submitForm": False,
                    "compactToolbar": False,
                    "toolbarNoTabs": False,
                    "hideRightMenu": False,
                    "about": False,
                    "feedback": False,
                    "goback": False
                },
                "coEditing": {
                    "mode": "fast",
                    "change": True
                }
            },
            "width": "100%",
            "height": "100%"
        }
        
        # JWT token处理
        if self.jwt_secret:
            try:
                import jwt
                payload = config.copy()
                payload["iat"] = int(time.time())
                payload["exp"] = int(time.time()) + 3600
                config["token"] = jwt.encode(payload, self.jwt_secret, algorithm='HS256')
            except Exception as e:
                self.logger.warning(f"JWT token生成失败: {e}")
        
        return config


    def register_routes(self, app: FastAPI):
        """注册OnlyOffice编辑器相关路由"""
        @app.get("/edit_url", response_class=HTMLResponse)
        async def edit_url_get(
            url: str = Query(..., description="文档URL"),
            filename: Optional[str] = Query(None, description="自定义文件名"),
            replace_information: Optional[str] = Query(None, description="替换规则JSON，格式：[{\"from\": \"原文本\", \"to\": \"新文本\"}]")
        ):
            """编辑在线文档 - 支持URL参数和内容替换"""
            try:
                if not url:
                    raise HTTPException(status_code=400, detail="缺少url参数")
                
                # 解析替换规则
                replace_rules = []
                if replace_information:
                    try:
                        replace_rules = json.loads(replace_information)
                        if not isinstance(replace_rules, list):
                            raise ValueError("replace_information必须是数组格式")
                        
                        # 验证替换规则格式
                        for rule in replace_rules:
                            if not isinstance(rule, dict) or 'from' not in rule or 'to' not in rule:
                                raise ValueError("替换规则必须包含from和to字段")
                        
                        self.logger.info(f"接收到替换规则: {replace_rules}")
                        
                    except json.JSONDecodeError as e:
                        raise HTTPException(status_code=400, detail=f"replace_information JSON格式错误: {str(e)}")
                    except ValueError as e:
                        raise HTTPException(status_code=400, detail=str(e))
                
                # 下载文件到临时目录
                file_path, auto_filename = self.download_file_from_url(url, filename)
                
                # 使用自定义filename（如果提供）
                final_filename = filename if filename else auto_filename
                
                # 执行内容替换（如果有替换规则）
                if replace_rules:
                    self.logger.info(f"开始执行内容替换，规则数量: {len(replace_rules)}")
                    file_path = self.replace_document_content(file_path, replace_rules)
                    self.logger.info("内容替换完成")
                
                # 生成OnlyOffice配置
                config = self.generate_onlyoffice_config(file_path, final_filename, url)
                
                # 构建替换信息显示文本
                replace_info_text = ""
                if replace_rules:
                    replace_info_text = f"<div style='font-size: 10px; margin-top: 5px; color: #666;'>已应用 {len(replace_rules)} 个替换规则</div>"
                
                # 生成HTML模板
                html_template = f'''
<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>编辑: {final_filename}</title>
    <style>
        body {{ margin: 0; padding: 0; overflow: hidden; height: 100vh; }}
        #editor {{ width: 100%; height: 100vh; border: none; }}
        .loading {{
            position: fixed; top: 50%; left: 50%; transform: translate(-50%, -50%);
            text-align: center; z-index: 1000; background: rgba(255,255,255,0.9);
            padding: 20px; border-radius: 10px; box-shadow: 0 4px 6px rgba(0,0,0,0.1);
        }}
        .spinner {{
            width: 40px; height: 40px; border: 4px solid #f3f3f3;
            border-top: 4px solid #4f46e5; border-radius: 50%;
            animation: spin 1s linear infinite; margin: 0 auto 10px;
        }}
        @keyframes spin {{ 0% {{ transform: rotate(0deg); }} 100% {{ transform: rotate(360deg); }} }}
    </style>
</head>
<body>
    <div class="loading" id="loading">
        <div class="spinner"></div>
        <div>正在加载编辑器...</div>
        <div style="font-size: 12px; margin-top: 10px;">{final_filename}</div>
        <div style="font-size: 10px; margin-top: 5px;">来源: {url}</div>
        {replace_info_text}
    </div>
    
    <div id="editor"></div>
    
    <script src="{self.onlyoffice_server}/web-apps/apps/api/documents/api.js"></script>
    <script>
        let config = {json.dumps(config)};
        
        console.log('文档URL:', '{url}');
        console.log('文件名:', '{final_filename}');
        console.log('替换规则:', {json.dumps(replace_rules)});
        console.log('表单权限:', config.document.permissions.fillForms);
        console.log('内容控件权限:', config.document.permissions.modifyContentControl);
        console.log('编辑模式:', config.editorConfig.mode);
        console.log('完整配置:', config);
        
        window.onload = function() {{
            try {{
                config.events = {{
                    'onAppReady': function() {{
                        console.log('编辑器已就绪，可以编辑表单字段');
                        document.getElementById('loading').style.display = 'none';
                    }},
                    'onDocumentReady': function() {{
                        console.log('文档已加载完成，表单字段应该可以编辑');
                    }},
                    'onRequestEditRights': function() {{
                        console.log('请求编辑权限');
                    }},
                    'onError': function(event) {{
                        console.error('编辑器错误:', event);
                        let errorMsg = '未知错误';
                        if (event && event.data) {{
                            switch(event.data) {{
                                case 1: errorMsg = '文档加载错误'; break;
                                case 2: errorMsg = '回调URL错误'; break;
                                case 3: errorMsg = '内部服务器错误'; break;
                                case 4: errorMsg = '文档密钥错误'; break;
                                case 5: errorMsg = '回调文档状态错误'; break;
                                case 6: errorMsg = '回调文档格式错误'; break;
                                default: errorMsg = `错误代码: ${{event.data}}`;
                            }}
                        }}
                        document.getElementById('loading').innerHTML = 
                            '<div>编辑器错误: ' + errorMsg + '</div>';
                    }}
                }};
                
                new DocsAPI.DocEditor("editor", config);
                console.log('编辑器实例已创建');
            }} catch(error) {{
                console.error('初始化失败:', error);
                document.getElementById('loading').innerHTML = 
                    '<div>加载失败: ' + error.message + '</div>';
            }}
        }};
    </script>
</body>
</html>
                '''
                
                return HTMLResponse(content=html_template, status_code=200)
                
            except HTTPException:
                raise
            except Exception as e:
                self.logger.error(f"编辑文档失败: {e}")
                raise HTTPException(status_code=500, detail=f"处理文档失败: {str(e)}")
        
        
        
        @app.post("/edit_url", response_class=HTMLResponse)
        async def edit_url_post(request: Request):
            """POST方式编辑在线文档 - 支持JSON请求体"""
            try:
                # 解析JSON请求体
                data = await request.json()
                
                url = data.get('url')
                filename = data.get('filename')
                replace_information = data.get('replace_information')
                
                if not url:
                    return JSONResponse(
                        status_code=400,
                        content={
                            "success": False,
                            "message": "url参数不能为空！",
                            "timestamp": datetime.now().isoformat()
                        }
                    )
                
                # 解析替换规则
                replace_rules = []
                if replace_information:
                    replace_rules = json.loads(replace_information) if isinstance(replace_information, str) else replace_information
                    
                    if not isinstance(replace_rules, list):
                        return JSONResponse(
                            status_code=400,
                            content={
                                "success": False,
                                "message": "replace_information必须是JSON格式",
                                "timestamp": datetime.now().isoformat()
                            }
                        )
                    
                    # 验证替换规则格式
                    for rule in replace_rules:
                        if not isinstance(rule, dict) or 'from' not in rule or 'to' not in rule:
                            return JSONResponse(
                                status_code=400,
                                content={
                                    "success": False,
                                    "message": "替换规则必须包含from和to字段",
                                    "timestamp": datetime.now().isoformat()
                                }
                            )
                    
                    self.logger.info(f"接收到替换规则: {replace_rules}")
                
                # 后续处理逻辑与GET请求相同
                file_path, auto_filename = self.download_file_from_url(url, filename)
                final_filename = filename if filename else auto_filename
                
                if replace_rules:
                    self.logger.info(f"开始执行内容替换，规则数量: {len(replace_rules)}")
                    file_path = self.replace_document_content(file_path, replace_rules)
                    self.logger.info("内容替换完成")
                
                config = self.generate_onlyoffice_config(file_path, final_filename, url)
                
                # 生成HTML响应（复用现有的HTML模板生成逻辑）
                # ... HTML模板代码与GET请求相同 ...
                # 构建替换信息显示文本
                replace_info_text = ""
                if replace_rules:
                    replace_info_text = f"<div style='font-size: 10px; margin-top: 5px; color: #666;'>已应用 {len(replace_rules)} 个替换规则</div>"
                
                # 生成HTML模板
                html_template = f'''
<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>编辑: {final_filename}</title>
    <style>
        body {{ margin: 0; padding: 0; overflow: hidden; height: 100vh; }}
        #editor {{ width: 100%; height: 100vh; border: none; }}
        .loading {{
            position: fixed; top: 50%; left: 50%; transform: translate(-50%, -50%);
            text-align: center; z-index: 1000; background: rgba(255,255,255,0.9);
            padding: 20px; border-radius: 10px; box-shadow: 0 4px 6px rgba(0,0,0,0.1);
        }}
        .spinner {{
            width: 40px; height: 40px; border: 4px solid #f3f3f3;
            border-top: 4px solid #4f46e5; border-radius: 50%;
            animation: spin 1s linear infinite; margin: 0 auto 10px;
        }}
        @keyframes spin {{ 0% {{ transform: rotate(0deg); }} 100% {{ transform: rotate(360deg); }} }}
    </style>
</head>
<body>
    <div class="loading" id="loading">
        <div class="spinner"></div>
        <div>正在加载编辑器...</div>
        <div style="font-size: 12px; margin-top: 10px;">{final_filename}</div>
        <div style="font-size: 10px; margin-top: 5px;">来源: {url}</div>
        {replace_info_text}
    </div>
    
    <div id="editor"></div>
    
    <script src="{self.onlyoffice_server}/web-apps/apps/api/documents/api.js"></script>
    <script>
        let config = {json.dumps(config)};
        
        console.log('文档URL:', '{url}');
        console.log('文件名:', '{final_filename}');
        console.log('替换规则:', {json.dumps(replace_rules)});
        console.log('表单权限:', config.document.permissions.fillForms);
        console.log('内容控件权限:', config.document.permissions.modifyContentControl);
        console.log('编辑模式:', config.editorConfig.mode);
        console.log('完整配置:', config);
        
        window.onload = function() {{
            try {{
                config.events = {{
                    'onAppReady': function() {{
                        console.log('编辑器已就绪，可以编辑表单字段');
                        document.getElementById('loading').style.display = 'none';
                    }},
                    'onDocumentReady': function() {{
                        console.log('文档已加载完成，表单字段应该可以编辑');
                    }},
                    'onRequestEditRights': function() {{
                        console.log('请求编辑权限');
                    }},
                    'onError': function(event) {{
                        console.error('编辑器错误:', event);
                        let errorMsg = '未知错误';
                        if (event && event.data) {{
                            switch(event.data) {{
                                case 1: errorMsg = '文档加载错误'; break;
                                case 2: errorMsg = '回调URL错误'; break;
                                case 3: errorMsg = '内部服务器错误'; break;
                                case 4: errorMsg = '文档密钥错误'; break;
                                case 5: errorMsg = '回调文档状态错误'; break;
                                case 6: errorMsg = '回调文档格式错误'; break;
                                default: errorMsg = `错误代码: ${{event.data}}`;
                            }}
                        }}
                        document.getElementById('loading').innerHTML = 
                            '<div>编辑器错误: ' + errorMsg + '</div>';
                    }}
                }};
                
                new DocsAPI.DocEditor("editor", config);
                console.log('编辑器实例已创建');
            }} catch(error) {{
                console.error('初始化失败:', error);
                document.getElementById('loading').innerHTML = 
                    '<div>加载失败: ' + error.message + '</div>';
            }}
        }};
    </script>
</body>
</html>
                '''
                return HTMLResponse(content=html_template, status_code=200)
            except Exception as e:
                return JSONResponse(
                    status_code=500,
                    content={
                        "success": False,
                        "message": f"处理文档失败: {str(e)}",
                        "timestamp": datetime.now().isoformat()
                    }
                )
        
        
        
        
        @app.get("/serve_temp_file")
        async def serve_temp_file(path: str = Query(..., description="临时文件路径")):
            """提供临时文件服务"""
            try:
                if not path:
                    raise HTTPException(status_code=400, detail="文件路径无效")
                
                # 解码路径
                file_path = urllib.parse.unquote(path)
                
                # 安全检查：确保文件在临时目录中
                if not file_path.startswith(self.temp_dir):
                    raise HTTPException(status_code=403, detail="文件路径无效")
                
                if not os.path.exists(file_path):
                    raise HTTPException(status_code=404, detail="文件不存在")
                
                # 获取文件扩展名来设置正确的Content-Type
                file_ext = os.path.splitext(file_path)[1].lower()
                content_type_mapping = {
                    '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    '.doc': 'application/msword',
                    '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                    '.xls': 'application/vnd.ms-excel',
                    '.pptx': 'application/vnd.openxmlformats-officedocument.presentationml.presentation',
                    '.ppt': 'application/vnd.ms-powerpoint'
                }
                content_type = content_type_mapping.get(file_ext, 'application/octet-stream')
                
                filename = os.path.basename(file_path)
                
                return FileResponse(
                    path=file_path,
                    media_type=content_type,
                    filename=filename
                )
            except HTTPException:
                raise
            except Exception as e:
                self.logger.error(f"文件服务错误: {e}")
                raise HTTPException(status_code=500, detail=f"文件服务错误: {str(e)}")
        
        
        @app.post("/callback")
        async def callback(
            request: Request,
            key: str = Query("unknown", description="文档密钥"),
            filename: str = Query("document.docx", description="文件名"),
            original_url: str = Query("", description="原始URL")
        ):
            """OnlyOffice回调 - 支持保存到API"""
            try:
                data = await request.json()
                
                # 解码filename
                filename = urllib.parse.unquote(filename)
                
                self.logger.info(f"OnlyOffice回调 - Key: {key}, Filename: {filename}, Status: {data.get('status')}")
                
                # 处理保存状态
                status = data.get('status', 0)
                if status == 2:
                    # 文档准备保存
                    download_url = data.get('url')
                    if download_url:
                        self.logger.info(f"文档需要保存，下载URL: {download_url}")
                        try:
                            # 下载编辑后的文档
                            response = requests.get(download_url, timeout=30)
                            response.raise_for_status()
                            
                            # 上传到保存API
                            upload_result = self.upload_file_to_api(response.content, filename)
                            self.logger.info(f"文档已保存到API: {upload_result}")
                            
                            # 清理临时文件
                            self.cleanup_temp_files()
                            
                        except Exception as e:
                            self.logger.error(f"保存文档时出错: {e}")
                            return JSONResponse(
                                status_code=200,
                                content={"error": 1, "message": str(e)}
                            )
                
                return JSONResponse(
                    status_code=200,
                    content={"error": 0}
                )
                
            except Exception as e:
                self.logger.error(f"回调处理错误: {e}")
                return JSONResponse(
                    status_code=200,
                    content={"error": 1, "message": str(e)}
                )
        
        
        @app.get("/health")
        async def health():
            """健康检查"""
            return JSONResponse(
                status_code=200,
                content={
                    "success": True,
                    "data": {
                        "status": "healthy",
                        "onlyoffice_server": self.onlyoffice_server,
                        "save_api_url": self.save_api_url,
                        "temp_dir": self.temp_dir
                    },
                    "timestamp": datetime.now().isoformat()
                }
            )


def create_app():
    app = FastAPI(
        title="OnlyOffice文档编辑器",
        description="支持URL传参编辑、内容替换和回调保存的OnlyOffice编辑器服务",
        version="2.1.0"
    )
    # 配置参数
    ONLYOFFICE_SERVER = 'http://ai.shunxikj.com:8442'
    SAVE_API_URL = 'https://ai.shunxikj.com:5002/api/files/upload'
    JWT_SECRET = ''  # 暂时禁用JWT

    editor = OnlyOfficeEditor(
        onlyoffice_server=ONLYOFFICE_SERVER,
        save_api_url=SAVE_API_URL,
        jwt_secret=JWT_SECRET
    )
    editor.register_routes(app)
    return app

            
if __name__ == '__main__':
    import uvicorn
    app = create_app()
    print("启动OnlyOffice编辑器: http://localhost:8002")
    print("支持URL传参编辑、内容替换和回调保存功能")
    print("保存API: https://ai.shunxikj.com:5002/api/files/upload")
    print("使用方式: /edit_url?url=文档URL&filename=文件名&replace_information=[{\"from\":\"原文本\",\"to\":\"新文本\"}]")
    
    uvicorn.run(app, host="0.0.0.0", port=8002, log_level="info")